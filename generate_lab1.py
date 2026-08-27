import os
import subprocess
import re
import shutil
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

CHROME_PATH = r"C:\Program Files\Google\Chrome\Application\chrome.exe"
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LAB1_DIR = os.path.join(BASE_DIR, "LAB1")
DOCS_DIR = os.path.join(LAB1_DIR, "docs")
DIAGRAM_DIR = os.path.join(LAB1_DIR, "diagram")

os.makedirs(LAB1_DIR, exist_ok=True)
os.makedirs(DOCS_DIR, exist_ok=True)
os.makedirs(DIAGRAM_DIR, exist_ok=True)

STUDENT_NAME = "Aditya T J"
STUDENT_SRN = "PES1UG24CS031"
PROBLEM_NUM = "31"
PROBLEM_TITLE = "Multi-Vendor Artisan E-Commerce Marketplace"
DOMAIN = "Retail, E-Commerce & Finance"

# ------------------------------------------------------------------------------
# Helper: Chrome Headless Print & Screenshot
# ------------------------------------------------------------------------------
def print_html_to_pdf(html_path, pdf_path, landscape=False):
    abs_html = os.path.abspath(html_path)
    abs_pdf = os.path.abspath(pdf_path)
    cmd = [
        CHROME_PATH,
        "--headless=new",
        "--disable-gpu",
        "--no-pdf-header-footer",
        f"--print-to-pdf={abs_pdf}",
        f"file:///{abs_html}"
    ]
    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode != 0:
        raise RuntimeError(f"Chrome failed to print {html_path}: {res.stderr}")
    print(f"  Generated PDF: {os.path.basename(pdf_path)} ({os.path.getsize(pdf_path)} bytes)")

def screenshot_html_to_png(html_path, png_path, width=1100, height=720):
    abs_html = os.path.abspath(html_path)
    abs_png = os.path.abspath(png_path)
    cmd = [
        CHROME_PATH,
        "--headless=new",
        "--disable-gpu",
        f"--window-size={width},{height}",
        f"--screenshot={abs_png}",
        f"file:///{abs_html}"
    ]
    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode != 0:
        raise RuntimeError(f"Chrome failed to screenshot {html_path}: {res.stderr}")
    print(f"  Generated PNG: {os.path.basename(png_path)} ({os.path.getsize(png_path)} bytes)")

def get_pdf_page_count(pdf_path):
    with open(pdf_path, 'rb') as f:
        data = f.read()
    return len(re.findall(rb'/Type\s*/Page\b', data))

# ==============================================================================
# DATA DEFINITIONS FOR REQUIREMENTS & FLOW
# ==============================================================================
REQUIREMENTS = [
    {
        "id": "FR-001",
        "type": "Functional",
        "desc": "The system shall split customer cart payments at checkout, allocating respective item earnings to multiple independent vendor accounts after deducting a 5% platform fee.",
        "priority": "High",
        "criteria": "Pass: Split payout calculations balance to total cart value.\nFail: Payout calculation discrepancies.",
        "rationale": "Ensures accurate, automated financial distribution to independent artisans while securing platform transaction revenue (Given in PS #31).",
        "comments": "Core monetization requirement; calculations must balance to 100% of order total."
    },
    {
        "id": "FR-002",
        "type": "Functional",
        "desc": "The system shall allow artisan vendors to set up and manage an independent storefront profile, including artisan bio, shop policies, and linked payout disbursement account details.",
        "priority": "High",
        "criteria": "Pass: Storefront profile updates and verified payout details are successfully saved and rendered publicly on the artisan's storefront page.\nFail: Missing mandatory profile fields or unvalidated payout details prevent storefront publishing.",
        "rationale": "Enables independent craftspeople to establish their distinct brand identity and receive automated financial payouts.",
        "comments": "Peer: Validated banking credentials required before activating vendor storefront."
    },
    {
        "id": "FR-003",
        "type": "Functional",
        "desc": "The system shall allow artisan vendors to create, update, and manage handcrafted product listings with titles, descriptions, pricing, inventory quantities, and high-resolution product media.",
        "priority": "High",
        "criteria": "Pass: New or updated product listings appear in the artisan's catalog and marketplace search within 5 seconds with accurate stock counts.\nFail: Listings with invalid pricing (<= 0) or missing mandatory attributes fail validation and are rejected.",
        "rationale": "Empowers artisans to maintain an up-to-date catalog of handcrafted items and prevent overselling through inventory synchronization.",
        "comments": "Real-time stock validation prevents race conditions during simultaneous customer checkouts."
    },
    {
        "id": "FR-004",
        "type": "Functional",
        "desc": "The system shall allow shoppers to add handcrafted items from multiple independent artisan vendors into a unified cart and execute a single consolidated checkout order.",
        "priority": "High",
        "criteria": "Pass: Cart aggregates items across distinct vendors, calculates itemized totals, applies applicable coupon discounts, and initiates payment authorization.\nFail: Cart fails to itemize multi-vendor items or calculation mismatch occurs between item sums and total payable amount.",
        "rationale": "Provides a seamless purchasing experience for shoppers buying from multiple artisans simultaneously without requiring separate checkout transactions.",
        "comments": "Consolidated order is decomposed into vendor sub-orders upon payment authorization."
    },
    {
        "id": "FR-005",
        "type": "Functional",
        "desc": "The system shall notify artisan vendors of confirmed customer orders containing their respective products and allow vendors to update fulfillment status (Processing, Dispatched, Delivered) with carrier tracking details.",
        "priority": "Medium",
        "criteria": "Pass: Vendor dashboard immediately displays incoming sub-orders upon successful checkout and status transitions trigger automated notifications to the shopper.\nFail: Order details fail to isolate vendor-specific items or status updates fail to persist.",
        "rationale": "Enables independent artisans to fulfill customer orders independently while keeping shoppers informed of dispatch status.",
        "comments": "Status changes send transactional push/email notifications with courier tracking links."
    },
    {
        "id": "NFR-001",
        "type": "Performance & Security",
        "desc": "The product catalog must support high-resolution image rendering with CDN caching delivering load times < 500 ms.",
        "priority": "High",
        "criteria": "Pass: Benchmarking tests confirm target latency (< 500 ms) and security standards under simulated peak load.\nFail: Catalog page load latency >= 500 ms or CDN caching failure under peak load.",
        "rationale": "Ensures fast page loading for image-heavy handcrafted goods, preserving user engagement and minimizing bounce rates (Given in PS #31).",
        "comments": "Edge CDN nodes cache optimized image formats (WebP/AVIF); benchmarked via Lighthouse."
    },
    {
        "id": "NFR-002",
        "type": "Security & Compliance",
        "desc": "The system shall encrypt all sensitive financial transactions and payout account data in transit using TLS 1.3 and at rest using AES-256 encryption, adhering to PCI-DSS Level 1 compliance standards.",
        "priority": "High",
        "criteria": "Pass: Automated vulnerability scans and penetration audits confirm 100% encryption coverage for payment payloads with zero plain-text storage of payment credentials.\nFail: Any unencrypted transmission or unmasked storage of sensitive banking/card data detected.",
        "rationale": "Protects shoppers' financial credentials and vendors' banking payout information against unauthorized access, data breaches, and regulatory penalties.",
        "comments": "Immutable append-only ledger verifies double-entry reconciliation between payouts and fees."
    }
]

# ==============================================================================
# 1. GENERATE REQUIREMENTS TABLE (PDF, DOCX, XLSX, HTML)
# ==============================================================================
def build_requirements_table():
    print("\n--- Building Requirements Table Deliverables ---")

    # 1A. HTML -> PDF
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>Requirements Table - Problem Statement #{PROBLEM_NUM}</title>
<style>
  @page {{
    size: A4 landscape;
    margin: 6mm 10mm 6mm 10mm;
  }}
  * {{
    box-sizing: border-box;
    margin: 0;
    padding: 0;
  }}
  body {{
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif;
    color: #111827;
    background: #ffffff;
    line-height: 1.24;
    font-size: 8pt;
  }}
  .header-card {{
    border: 1.5px solid #1e3a8a;
    border-radius: 5px;
    padding: 6px 12px;
    margin-bottom: 5px;
    background: linear-gradient(135deg, #f0f4f8 0%, #ffffff 100%);
    display: flex;
    justify-content: space-between;
    align-items: center;
  }}
  .header-left h1 {{
    font-size: 11.5pt;
    color: #1e3a8a;
    font-weight: 700;
    letter-spacing: -0.2px;
    margin-bottom: 2px;
  }}
  .header-left h2 {{
    font-size: 8.5pt;
    color: #374151;
    font-weight: 600;
  }}
  .header-right {{
    text-align: right;
    font-size: 7.8pt;
    color: #1f2937;
  }}
  .badge-tag {{
    display: inline-block;
    background: #1e3a8a;
    color: white;
    padding: 2px 8px;
    border-radius: 4px;
    font-weight: bold;
    font-size: 7.4pt;
    letter-spacing: 0.3px;
  }}
  .context-box {{
    background: #f8fafc;
    border-left: 3.5px solid #2563eb;
    padding: 4px 10px;
    margin-bottom: 5px;
    font-size: 7.6pt;
    color: #334155;
    border-radius: 0 4px 4px 0;
  }}
  table {{
    width: 100%;
    border-collapse: collapse;
  }}
  th {{
    background-color: #1e3a8a;
    color: #ffffff;
    font-weight: 600;
    text-align: left;
    padding: 4.5px 6px;
    font-size: 7.6pt;
    border: 1px solid #1e3a8a;
    text-transform: uppercase;
    letter-spacing: 0.3px;
  }}
  td {{
    padding: 3.5px 6px;
    border: 1px solid #cbd5e1;
    vertical-align: top;
    font-size: 7.3pt;
    color: #1f2937;
  }}
  tr:nth-child(even) {{
    background-color: #f8fafc;
  }}
  .req-id {{
    font-weight: 700;
    color: #1e3a8a;
    white-space: nowrap;
    text-align: center;
  }}
  .badge-func {{
    display: inline-block;
    padding: 1.5px 5px;
    border-radius: 3px;
    font-weight: 600;
    font-size: 6.8pt;
    background: #e0e7ff;
    color: #3730a3;
  }}
  .badge-nonfunc {{
    display: inline-block;
    padding: 1.5px 5px;
    border-radius: 3px;
    font-weight: 600;
    font-size: 6.8pt;
    background: #fef3c7;
    color: #92400e;
  }}
  .badge-high {{
    display: inline-block;
    padding: 1.5px 5px;
    border-radius: 3px;
    font-weight: 600;
    font-size: 6.8pt;
    background: #fee2e2;
    color: #991b1b;
  }}
  .badge-med {{
    display: inline-block;
    padding: 1.5px 5px;
    border-radius: 3px;
    font-weight: 600;
    font-size: 6.8pt;
    background: #e0f2fe;
    color: #075985;
  }}
  .pass-text {{
    color: #15803d;
    font-weight: 700;
  }}
  .fail-text {{
    color: #b91c1c;
    font-weight: 700;
  }}
  .section-divider {{
    background-color: #e2e8f0;
    font-weight: 700;
    color: #0f172a;
    padding: 3px 6px;
    text-transform: uppercase;
    letter-spacing: 0.5px;
    font-size: 7.2pt;
  }}
  .footer {{
    margin-top: 5px;
    display: flex;
    justify-content: space-between;
    font-size: 6.8pt;
    color: #64748b;
    border-top: 1px solid #e2e8f0;
    padding-top: 3px;
  }}
</style>
</head>
<body>

<div class="header-card">
  <div class="header-left">
    <h1>PES UNIVERSITY &bull; DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING</h1>
    <h2>Software Engineering Lab 1: Requirements Engineering & UML Use-Case Modelling</h2>
  </div>
  <div class="header-right">
    <div><span class="badge-tag">PROBLEM STATEMENT #{PROBLEM_NUM}</span></div>
    <div style="margin-top: 3px;"><strong>Student:</strong> {STUDENT_NAME} ({STUDENT_SRN})</div>
    <div><strong>Domain:</strong> {DOMAIN}</div>
  </div>
</div>

<div class="context-box">
  <strong>System Overview:</strong> An online marketplace enabling independent craftspeople to set up storefronts, manage product catalogs, receive orders, and receive automated split payouts with platform commission deductions. 
  &bull; <strong>Target Actors:</strong> Shopper, Artisan Vendor, Payment Gateway (External System)
</div>

<table>
  <thead>
    <tr>
      <th style="width: 6.5%; text-align: center;">Req ID</th>
      <th style="width: 11%;">Type</th>
      <th style="width: 25%;">Description</th>
      <th style="width: 6.5%; text-align: center;">Priority</th>
      <th style="width: 25%;">Acceptance Criteria</th>
      <th style="width: 14%;">Rationale</th>
      <th style="width: 12%;">Comments</th>
    </tr>
  </thead>
  <tbody>
    <tr>
      <td colspan="7" class="section-divider">Functional Requirements (FR-001 to FR-005)</td>
    </tr>
"""
    for req in REQUIREMENTS[:5]:
        p_badge = 'badge-high' if req['priority'] == 'High' else 'badge-med'
        c_fmt = req['criteria'].replace('Pass:', '<span class="pass-text">Pass:</span>').replace('Fail:', '<span class="fail-text">Fail:</span>').replace('\n', '<br>')
        html += f"""    <tr>
      <td class="req-id">{req['id']}</td>
      <td><span class="badge-func">{req['type']}</span></td>
      <td>{req['desc']}</td>
      <td style="text-align: center;"><span class="{p_badge}">{req['priority']}</span></td>
      <td>{c_fmt}</td>
      <td>{req['rationale']}</td>
      <td>{req['comments']}</td>
    </tr>\n"""

    html += """    <tr>
      <td colspan="7" class="section-divider">Non-Functional Requirements (NFR-001 & NFR-002)</td>
    </tr>\n"""

    for req in REQUIREMENTS[5:]:
        p_badge = 'badge-high' if req['priority'] == 'High' else 'badge-med'
        c_fmt = req['criteria'].replace('Pass:', '<span class="pass-text">Pass:</span>').replace('Fail:', '<span class="fail-text">Fail:</span>').replace('\n', '<br>')
        html += f"""    <tr>
      <td class="req-id">{req['id']}</td>
      <td><span class="badge-nonfunc">{req['type']}</span></td>
      <td>{req['desc']}</td>
      <td style="text-align: center;"><span class="{p_badge}">{req['priority']}</span></td>
      <td>{c_fmt}</td>
      <td>{req['rationale']}</td>
      <td>{req['comments']}</td>
    </tr>\n"""

    html += f"""  </tbody>
</table>

<div class="footer">
  <span>Software Engineering Lab (UE24CS252) &bull; PES University Department of Computer Science & Engineering</span>
  <span>Student: {STUDENT_NAME} ({STUDENT_SRN})</span>
  <span>Problem Statement #{PROBLEM_NUM}: {PROBLEM_TITLE} &bull; Page 1 of 1</span>
</div>

</body>
</html>
"""
    html_path = os.path.join(DOCS_DIR, "01_Requirements_Table.html")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html)
    
    pdf_docs_path = os.path.join(DOCS_DIR, "01_Requirements_Table.pdf")
    pdf_root_path = os.path.join(LAB1_DIR, "Requirements_Table.pdf")
    print_html_to_pdf(html_path, pdf_docs_path, landscape=True)
    shutil.copyfile(pdf_docs_path, pdf_root_path)

    # 1B. DOCX (Word Document)
    docx_path = os.path.join(DOCS_DIR, "01_Requirements_Table.docx")
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)

    title_p = doc.add_paragraph()
    title_run = title_p.add_run("PES UNIVERSITY • DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING\n")
    title_run.bold = True
    title_run.font.size = Pt(13)
    title_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    sub_run = title_p.add_run(f"Software Engineering Lab 1: Requirements Engineering | Problem Statement #{PROBLEM_NUM}: {PROBLEM_TITLE}\nStudent: {STUDENT_NAME} ({STUDENT_SRN}) • Domain: {DOMAIN}")
    sub_run.font.size = Pt(9.5)
    sub_run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["Req ID", "Type", "Description", "Priority", "Acceptance Criteria", "Rationale", "Comments"]
    widths = [Inches(0.8), Inches(1.2), Inches(3.2), Inches(0.8), Inches(2.8), Inches(1.5), Inches(1.4)]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].width = widths[i]
        shading = parse_xml(r'<w:shd {} w:fill="1E3A8A"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for req in REQUIREMENTS:
        row_cells = table.add_row().cells
        data = [req["id"], req["type"], req["desc"], req["priority"], req["criteria"], req["rationale"], req["comments"]]
        for i, val in enumerate(data):
            row_cells[i].text = val
            row_cells[i].width = widths[i]
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
                    r.font.name = "Segoe UI"
    doc.save(docx_path)
    print(f"  Generated DOCX: {os.path.basename(docx_path)}")

    # 1C. XLSX (Excel Spreadsheet)
    xlsx_path = os.path.join(DOCS_DIR, "01_Requirements_Table.xlsx")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Requirements Table"

    ws.merge_cells("A1:G1")
    ws["A1"] = f"PES UNIVERSITY - CSE • Software Engineering Lab 1: Requirements Table (PS #{PROBLEM_NUM})"
    ws["A1"].font = Font(name="Segoe UI", size=12, bold=True, color="FFFFFF")
    ws["A1"].fill = PatternFill(start_color="1E3A8A", end_color="1E3A8A", fill_type="solid")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")

    ws.merge_cells("A2:G2")
    ws["A2"] = f"Student: {STUDENT_NAME} ({STUDENT_SRN}) | Topic: {PROBLEM_TITLE} | Domain: {DOMAIN}"
    ws["A2"].font = Font(name="Segoe UI", size=9.5, italic=True, color="1E3A8A")
    ws["A2"].alignment = Alignment(horizontal="center", vertical="center")

    hdr_row = ["Req ID", "Type", "Description", "Priority", "Acceptance Criteria", "Rationale", "Comments"]
    ws.append(hdr_row)
    for col_num in range(1, 8):
        c = ws.cell(row=3, column=col_num)
        c.font = Font(name="Segoe UI", size=10, bold=True, color="FFFFFF")
        c.fill = PatternFill(start_color="2563EB", end_color="2563EB", fill_type="solid")
        c.alignment = Alignment(horizontal="center" if col_num in [1, 4] else "left", vertical="center")

    thin_border = Border(
        left=Side(style='thin', color='CBD5E1'),
        right=Side(style='thin', color='CBD5E1'),
        top=Side(style='thin', color='CBD5E1'),
        bottom=Side(style='thin', color='CBD5E1')
    )

    for req in REQUIREMENTS:
        ws.append([req["id"], req["type"], req["desc"], req["priority"], req["criteria"], req["rationale"], req["comments"]])
        row_idx = ws.max_row
        fill = PatternFill(start_color="F8FAFC" if row_idx % 2 == 0 else "FFFFFF", fill_type="solid")
        for col_num in range(1, 8):
            cell = ws.cell(row=row_idx, column=col_num)
            cell.font = Font(name="Segoe UI", size=9)
            cell.border = thin_border
            cell.fill = fill
            cell.alignment = Alignment(wrap_text=True, vertical="top", horizontal="center" if col_num in [1, 4] else "left")

    col_widths = [12, 18, 48, 12, 45, 32, 28]
    for i, w in enumerate(col_widths):
        ws.column_dimensions[get_column_letter(i+1)].width = w

    wb.save(xlsx_path)
    print(f"  Generated XLSX: {os.path.basename(xlsx_path)}")

    shutil.copyfile(docx_path, os.path.join(LAB1_DIR, "Requirements_Table.docx"))
    shutil.copyfile(xlsx_path, os.path.join(LAB1_DIR, "Requirements_Table.xlsx"))

# ==============================================================================
# 2. GENERATE USE-CASE DIAGRAM (SVG, DRAWIO, PNG, PDF)
# ==============================================================================
def build_use_case_diagram():
    print("\n--- Building UML Use-Case Diagram Deliverables ---")

    svg_content = f"""<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 1060 680" width="1060" height="680" style="background:#ffffff; font-family:-apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Arial, sans-serif;">
  <defs>
    <!-- Marker for dashed include/extend arrows -->
    <marker id="arrow" viewBox="0 0 10 10" refX="9" refY="5" markerWidth="8" markerHeight="8" orient="auto-start-reverse">
      <path d="M 0 1 L 10 5 L 0 9 z" fill="#1e3a8a" />
    </marker>
    <!-- Drop shadow filter -->
    <filter id="shadow" x="-5%" y="-5%" width="110%" height="115%" filterUnits="userSpaceOnUse">
      <feDropShadow dx="2" dy="2" stdDeviation="2.5" flood-opacity="0.12" />
    </filter>
  </defs>

  <!-- Title & Header Banner -->
  <rect x="25" y="15" width="1010" height="52" rx="6" fill="#f0f4f8" stroke="#1e3a8a" stroke-width="1.5"/>
  <text x="45" y="38" font-size="16" font-weight="bold" fill="#1e3a8a">UML USE-CASE DIAGRAM: Multi-Vendor Artisan E-Commerce Marketplace</text>
  <text x="45" y="55" font-size="11" fill="#4b5563">PES University Dept. of CSE &bull; Problem Statement #{PROBLEM_NUM} &bull; Student: {STUDENT_NAME} ({STUDENT_SRN})</text>
  <rect x="880" y="27" width="135" height="26" rx="4" fill="#1e3a8a"/>
  <text x="947" y="44" font-size="11" font-weight="bold" fill="#ffffff" text-anchor="middle">RETAIL &amp; E-COMMERCE</text>

  <!-- SYSTEM BOUNDARY -->
  <rect x="230" y="85" width="580" height="570" rx="10" fill="#fbfcfe" stroke="#1e3a8a" stroke-width="2" stroke-dasharray="6,4"/>
  <!-- Boundary Header Box -->
  <rect x="230" y="85" width="580" height="34" rx="10" fill="#1e3a8a"/>
  <rect x="230" y="105" width="580" height="14" fill="#1e3a8a"/>
  <text x="520" y="107" font-size="13" font-weight="bold" fill="#ffffff" text-anchor="middle" letter-spacing="0.5">Multi-Vendor Artisan E-Commerce Marketplace (System Boundary)</text>

  <!-- ================= ACTOR 1: SHOPPER (LEFT TOP) ================= -->
  <g id="actor-shopper">
    <!-- Stick figure -->
    <circle cx="105" cy="180" r="18" fill="#e0e7ff" stroke="#1e3a8a" stroke-width="2"/>
    <line x1="105" y1="198" x2="105" y2="245" stroke="#1e3a8a" stroke-width="2.2"/>
    <line x1="75" y1="215" x2="135" y2="215" stroke="#1e3a8a" stroke-width="2.2"/>
    <line x1="105" y1="245" x2="80" y2="285" stroke="#1e3a8a" stroke-width="2.2"/>
    <line x1="105" y1="245" x2="130" y2="285" stroke="#1e3a8a" stroke-width="2.2"/>
    <!-- Actor Box/Label -->
    <rect x="50" y="295" width="110" height="30" rx="4" fill="#1e3a8a"/>
    <text x="105" y="314" font-size="12" font-weight="bold" fill="#ffffff" text-anchor="middle">Shopper</text>
    <text x="105" y="340" font-size="10" fill="#4b5563" text-anchor="middle">(Primary Actor)</text>
  </g>

  <!-- ================= ACTOR 2: ARTISAN VENDOR (LEFT BOTTOM) ================= -->
  <g id="actor-artisan">
    <!-- Stick figure -->
    <circle cx="105" cy="460" r="18" fill="#fef3c7" stroke="#92400e" stroke-width="2"/>
    <line x1="105" y1="478" x2="105" y2="525" stroke="#92400e" stroke-width="2.2"/>
    <line x1="75" y1="495" x2="135" y2="495" stroke="#92400e" stroke-width="2.2"/>
    <line x1="105" y1="525" x2="80" y2="565" stroke="#92400e" stroke-width="2.2"/>
    <line x1="105" y1="525" x2="130" y2="565" stroke="#92400e" stroke-width="2.2"/>
    <!-- Actor Box/Label -->
    <rect x="40" y="575" width="130" height="30" rx="4" fill="#92400e"/>
    <text x="105" y="594" font-size="12" font-weight="bold" fill="#ffffff" text-anchor="middle">Artisan Vendor</text>
    <text x="105" y="620" font-size="10" fill="#4b5563" text-anchor="middle">(Primary Actor)</text>
  </g>

  <!-- ================= ACTOR 3: PAYMENT GATEWAY (RIGHT) ================= -->
  <g id="actor-payment-gateway">
    <!-- System Actor Rectangular Icon & Stick Figure -->
    <rect x="870" y="270" width="145" height="75" rx="6" fill="#f0fdf4" stroke="#15803d" stroke-width="2" filter="url(#shadow)"/>
    <rect x="880" y="278" width="125" height="20" rx="3" fill="#15803d"/>
    <text x="942" y="292" font-size="10" font-weight="bold" fill="#ffffff" text-anchor="middle">&lt;&lt;external system&gt;&gt;</text>
    <text x="942" y="318" font-size="13" font-weight="bold" fill="#166534" text-anchor="middle">Payment</text>
    <text x="942" y="335" font-size="13" font-weight="bold" fill="#166534" text-anchor="middle">Gateway</text>
    <text x="942" y="365" font-size="10" fill="#4b5563" text-anchor="middle">(Secondary Actor)</text>
  </g>

  <!-- ================= USE CASES (OVALS INSIDE BOUNDARY) ================= -->
  <!-- UC-01: Browse & Search Catalog -->
  <g id="uc-01">
    <ellipse cx="375" cy="160" rx="110" ry="24" fill="#ffffff" stroke="#1e3a8a" stroke-width="2" filter="url(#shadow)"/>
    <text x="375" y="164" font-size="11" font-weight="bold" fill="#1e3a8a" text-anchor="middle">UC-01: Browse &amp; Search Catalog</text>
  </g>

  <!-- UC-02: Manage Shopping Cart -->
  <g id="uc-02">
    <ellipse cx="375" cy="225" rx="110" ry="24" fill="#ffffff" stroke="#1e3a8a" stroke-width="2" filter="url(#shadow)"/>
    <text x="375" y="229" font-size="11" font-weight="bold" fill="#1e3a8a" text-anchor="middle">UC-02: Manage Shopping Cart</text>
  </g>

  <!-- UC-03: Place Order (Core Centerpiece) -->
  <g id="uc-03">
    <ellipse cx="375" cy="305" rx="115" ry="26" fill="#eff6ff" stroke="#1d4ed8" stroke-width="2.5" filter="url(#shadow)"/>
    <text x="375" y="303" font-size="12" font-weight="bold" fill="#1e3a8a" text-anchor="middle">UC-03: Place Order</text>
    <text x="375" y="318" font-size="9" fill="#2563eb" text-anchor="middle">[Multi-Vendor Checkout]</text>
  </g>

  <!-- UC-05: Apply Discount Code (Extends UC-03) -->
  <g id="uc-05">
    <ellipse cx="375" cy="400" rx="110" ry="24" fill="#ffffff" stroke="#1e3a8a" stroke-width="2" filter="url(#shadow)"/>
    <text x="375" y="404" font-size="11" font-weight="bold" fill="#1e3a8a" text-anchor="middle">UC-05: Apply Discount Code</text>
  </g>

  <!-- UC-04: Process Split Payment (Included by UC-03) -->
  <g id="uc-04">
    <ellipse cx="675" cy="305" rx="118" ry="28" fill="#f0fdf4" stroke="#15803d" stroke-width="2.5" filter="url(#shadow)"/>
    <text x="675" y="299" font-size="11.5" font-weight="bold" fill="#166534" text-anchor="middle">UC-04: Process Split Payment</text>
    <text x="675" y="314" font-size="9" fill="#15803d" text-anchor="middle">[5% Fee &amp; Vendor Payouts]</text>
  </g>

  <!-- UC-06: Manage Storefront & Catalog -->
  <g id="uc-06">
    <ellipse cx="430" cy="505" rx="135" ry="25" fill="#ffffff" stroke="#92400e" stroke-width="2" filter="url(#shadow)"/>
    <text x="430" y="509" font-size="11" font-weight="bold" fill="#92400e" text-anchor="middle">UC-06: Manage Storefront &amp; Catalog</text>
  </g>

  <!-- UC-07: Fulfill Order -->
  <g id="uc-07">
    <ellipse cx="430" cy="580" rx="115" ry="24" fill="#ffffff" stroke="#92400e" stroke-width="2" filter="url(#shadow)"/>
    <text x="430" y="584" font-size="11" font-weight="bold" fill="#92400e" text-anchor="middle">UC-07: Fulfill Order</text>
  </g>

  <!-- ================= ASSOCIATIONS (SOLID LINES) ================= -->
  <!-- Shopper to UC-01 -->
  <line x1="160" y1="210" x2="268" y2="168" stroke="#1e3a8a" stroke-width="1.8"/>
  <!-- Shopper to UC-02 -->
  <line x1="160" y1="230" x2="266" y2="225" stroke="#1e3a8a" stroke-width="1.8"/>
  <!-- Shopper to UC-03 -->
  <line x1="160" y1="250" x2="265" y2="295" stroke="#1e3a8a" stroke-width="1.8"/>

  <!-- Artisan Vendor to UC-06 -->
  <line x1="170" y1="520" x2="298" y2="510" stroke="#92400e" stroke-width="1.8"/>
  <!-- Artisan Vendor to UC-07 -->
  <line x1="170" y1="540" x2="318" y2="575" stroke="#92400e" stroke-width="1.8"/>

  <!-- Payment Gateway to UC-04 -->
  <line x1="870" y1="305" x2="793" y2="305" stroke="#15803d" stroke-width="1.8"/>

  <!-- ================= STEREOTYPED RELATIONSHIPS ================= -->
  <!-- 1. <<include>> : UC-03 (Place Order) to UC-04 (Process Split Payment) -->
  <line x1="490" y1="305" x2="550" y2="305" stroke="#1e3a8a" stroke-width="2" stroke-dasharray="6,4" marker-end="url(#arrow)"/>
  <rect x="495" y="280" width="56" height="18" rx="3" fill="#ffffff" stroke="#cbd5e1" stroke-width="1"/>
  <text x="523" y="293" font-size="9" font-weight="bold" fill="#1e3a8a" text-anchor="middle">&lt;&lt;include&gt;&gt;</text>

  <!-- 2. <<extend>> : UC-05 (Apply Discount Code) to UC-03 (Place Order) -->
  <line x1="375" y1="376" x2="375" y2="339" stroke="#1e3a8a" stroke-width="2" stroke-dasharray="6,4" marker-end="url(#arrow)"/>
  <rect x="345" y="347" width="60" height="18" rx="3" fill="#ffffff" stroke="#cbd5e1" stroke-width="1"/>
  <text x="375" y="360" font-size="9" font-weight="bold" fill="#1e3a8a" text-anchor="middle">&lt;&lt;extend&gt;&gt;</text>

  <!-- ================= LEGEND CARD ================= -->
  <rect x="710" y="470" width="295" height="175" rx="6" fill="#f8fafc" stroke="#94a3b8" stroke-width="1.2" filter="url(#shadow)"/>
  <rect x="710" y="470" width="295" height="24" rx="6" fill="#e2e8f0"/>
  <rect x="710" y="488" width="295" height="6" fill="#e2e8f0"/>
  <text x="725" y="486" font-size="11" font-weight="bold" fill="#1e293b">UML Notation &amp; Relationship Legend</text>

  <!-- Legend item 1: Association -->
  <line x1="725" y1="512" x2="775" y2="512" stroke="#1e3a8a" stroke-width="1.8"/>
  <text x="785" y="516" font-size="10" fill="#334155"><tspan font-weight="bold">Association:</tspan> Direct actor-system interaction</text>

  <!-- Legend item 2: <<include>> -->
  <line x1="725" y1="542" x2="767" y2="542" stroke="#1e3a8a" stroke-width="1.8" stroke-dasharray="5,3" marker-end="url(#arrow)"/>
  <text x="785" y="546" font-size="10" fill="#334155"><tspan font-weight="bold">&lt;&lt;include&gt;&gt;:</tspan> Mandatory sub-routine execution</text>

  <!-- Legend item 3: <<extend>> -->
  <line x1="725" y1="572" x2="767" y2="572" stroke="#1e3a8a" stroke-width="1.8" stroke-dasharray="5,3" marker-end="url(#arrow)"/>
  <text x="785" y="576" font-size="10" fill="#334155"><tspan font-weight="bold">&lt;&lt;extend&gt;&gt;:</tspan> Optional / conditional enhancement</text>

  <!-- Legend item 4: System Boundary -->
  <rect x="725" y="596" width="45" height="15" rx="3" fill="#fbfcfe" stroke="#1e3a8a" stroke-width="1.5" stroke-dasharray="4,2"/>
  <text x="785" y="608" font-size="10" fill="#334155"><tspan font-weight="bold">System Boundary:</tspan> Scope of Marketplace</text>

  <!-- Legend item 5: 5% Commission Note -->
  <text x="725" y="632" font-size="9.5" fill="#15803d" font-weight="600">&bull; FR-001 Aligned: Automated 5% fee split deduction</text>
</svg>"""

    # 2A. Save pure SVG to diagram/ and root
    svg_diag_path = os.path.join(DIAGRAM_DIR, "artisan_marketplace_usecase.svg")
    with open(svg_diag_path, "w", encoding="utf-8") as f:
        f.write(svg_content)
    shutil.copyfile(svg_diag_path, os.path.join(LAB1_DIR, "Use_Case_Diagram.svg"))
    print(f"  Generated SVG: {os.path.basename(svg_diag_path)}")

    # 2B. Wrap in HTML for PDF export & Screenshot PNG
    html_page = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>Use Case Diagram - Problem Statement #{PROBLEM_NUM}</title>
<style>
  @page {{
    size: A4 landscape;
    margin: 6mm;
  }}
  * {{
    box-sizing: border-box;
    margin: 0;
    padding: 0;
  }}
  body {{
    display: flex;
    justify-content: center;
    align-items: center;
    height: 100vh;
    background: #ffffff;
    overflow: hidden;
  }}
  svg {{
    width: 100%;
    max-height: 98vh;
  }}
</style>
</head>
<body>
{svg_content}
</body>
</html>"""
    html_path = os.path.join(DOCS_DIR, "02_UseCase_Diagram.html")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_page)

    pdf_docs_path = os.path.join(DOCS_DIR, "02_UseCase_Diagram.pdf")
    pdf_root_path = os.path.join(LAB1_DIR, "Use_Case_Diagram.pdf")
    print_html_to_pdf(html_path, pdf_docs_path, landscape=True)
    shutil.copyfile(pdf_docs_path, pdf_root_path)
    shutil.copyfile(pdf_docs_path, os.path.join(LAB1_DIR, "UseCase_Diagram.pdf"))

    # 2C. Screenshot HTML to high-res PNG
    png_docs_path = os.path.join(DOCS_DIR, "02_UseCase_Diagram.png")
    png_root_path = os.path.join(LAB1_DIR, "use_case_diagram.png")
    screenshot_html_to_png(html_path, png_docs_path, width=1080, height=700)
    shutil.copyfile(png_docs_path, png_root_path)
    shutil.copyfile(png_docs_path, os.path.join(LAB1_DIR, "UseCase_Diagram.png"))
    shutil.copyfile(png_docs_path, os.path.join(LAB1_DIR, "02_UseCase_Diagram.png"))

    # 2D. Draw.io editable XML
    drawio_xml = """<mxfile host="app.diagrams.net" modified="2024-08-27T17:00:00.000Z" agent="Antigravity SE Lab 1" version="22.1.18" type="device">
  <diagram id="multi-vendor-marketplace" name="Multi-Vendor Artisan E-Commerce Marketplace">
    <mxGraphModel dx="1200" dy="800" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="1169" pageHeight="827" background="#ffffff" math="0" shadow="0">
      <root>
        <mxCell id="0" />
        <mxCell id="1" parent="0" />
        
        <!-- System Boundary -->
        <mxCell id="boundary" value="Multi-Vendor Artisan E-Commerce Marketplace" style="shape=swimlane;whiteSpace=wrap;html=1;startSize=30;horizontal=1;container=1;collapsible=0;strokeWidth=2;fillColor=#F8FAFC;strokeColor=#1E3A8A;fontStyle=1;fontSize=14;fontColor=#1E3A8A;" vertex="1" parent="1">
          <mxGeometry x="260" y="80" width="580" height="640" as="geometry" />
        </mxCell>
        
        <!-- Use Cases inside boundary -->
        <mxCell id="uc1" value="UC-01: Browse &amp; Search Catalog" style="ellipse;whiteSpace=wrap;html=1;fillColor=#FFFFFF;strokeColor=#1E3A8A;strokeWidth=2;fontSize=12;fontStyle=1;" vertex="1" parent="boundary">
          <mxGeometry x="40" y="60" width="220" height="50" as="geometry" />
        </mxCell>
        
        <mxCell id="uc2" value="UC-02: Manage Shopping Cart" style="ellipse;whiteSpace=wrap;html=1;fillColor=#FFFFFF;strokeColor=#1E3A8A;strokeWidth=2;fontSize=12;fontStyle=1;" vertex="1" parent="boundary">
          <mxGeometry x="40" y="140" width="220" height="50" as="geometry" />
        </mxCell>
        
        <mxCell id="uc3" value="UC-03: Place Order&#xa;[Multi-Vendor Checkout]" style="ellipse;whiteSpace=wrap;html=1;fillColor=#EFF6FF;strokeColor=#1D4ED8;strokeWidth=2.5;fontSize=12;fontStyle=1;fontColor=#1E3A8A;" vertex="1" parent="boundary">
          <mxGeometry x="40" y="230" width="230" height="60" as="geometry" />
        </mxCell>
        
        <mxCell id="uc5" value="UC-05: Apply Discount Code" style="ellipse;whiteSpace=wrap;html=1;fillColor=#FFFFFF;strokeColor=#1E3A8A;strokeWidth=2;fontSize=12;fontStyle=1;" vertex="1" parent="boundary">
          <mxGeometry x="45" y="340" width="220" height="50" as="geometry" />
        </mxCell>
        
        <mxCell id="uc4" value="UC-04: Process Split Payment&#xa;[5% Fee &amp; Vendor Escrow]" style="ellipse;whiteSpace=wrap;html=1;fillColor=#F0FDF4;strokeColor=#15803D;strokeWidth=2.5;fontSize=12;fontStyle=1;fontColor=#166534;" vertex="1" parent="boundary">
          <mxGeometry x="320" y="230" width="230" height="60" as="geometry" />
        </mxCell>
        
        <mxCell id="uc6" value="UC-06: Manage Storefront &amp; Catalog" style="ellipse;whiteSpace=wrap;html=1;fillColor=#FFFFFF;strokeColor=#92400E;strokeWidth=2;fontSize=12;fontStyle=1;fontColor=#92400E;" vertex="1" parent="boundary">
          <mxGeometry x="120" y="450" width="260" height="55" as="geometry" />
        </mxCell>
        
        <mxCell id="uc7" value="UC-07: Fulfill Order" style="ellipse;whiteSpace=wrap;html=1;fillColor=#FFFFFF;strokeColor=#92400E;strokeWidth=2;fontSize=12;fontStyle=1;fontColor=#92400E;" vertex="1" parent="boundary">
          <mxGeometry x="140" y="540" width="220" height="50" as="geometry" />
        </mxCell>
        
        <!-- Include relationship: Place Order -> Process Split Payment -->
        <mxCell id="rel_inc" value="&amp;lt;&amp;lt;include&amp;gt;&amp;gt;" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;dashed=1;endArrow=open;endFill=0;strokeWidth=2;strokeColor=#1E3A8A;fontStyle=1;fontSize=11;labelBackgroundColor=#FFFFFF;" edge="1" parent="boundary" source="uc3" target="uc4">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        
        <!-- Extend relationship: Apply Discount Code -> Place Order -->
        <mxCell id="rel_ext" value="&amp;lt;&amp;lt;extend&amp;gt;&amp;gt;" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;dashed=1;endArrow=open;endFill=0;strokeWidth=2;strokeColor=#1E3A8A;fontStyle=1;fontSize=11;labelBackgroundColor=#FFFFFF;" edge="1" parent="boundary" source="uc5" target="uc3">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>

        <!-- Actors outside boundary -->
        <mxCell id="act_shopper" value="Shopper&#xa;(Primary Actor)" style="shape=umlActor;verticalLabelPosition=bottom;verticalAlign=top;html=1;outlineConnect=0;fillColor=#E0E7FF;strokeColor=#1E3A8A;strokeWidth=2;fontStyle=1;fontSize=12;" vertex="1" parent="1">
          <mxGeometry x="90" y="210" width="60" height="110" as="geometry" />
        </mxCell>
        
        <mxCell id="act_artisan" value="Artisan Vendor&#xa;(Primary Actor)" style="shape=umlActor;verticalLabelPosition=bottom;verticalAlign=top;html=1;outlineConnect=0;fillColor=#FEF3C7;strokeColor=#92400E;strokeWidth=2;fontStyle=1;fontSize=12;" vertex="1" parent="1">
          <mxGeometry x="90" y="520" width="60" height="110" as="geometry" />
        </mxCell>
        
        <mxCell id="act_gateway" value="&amp;lt;&amp;lt;external system&amp;gt;&amp;gt;&#xa;Payment Gateway&#xa;(Secondary Actor)" style="shape=umlActor;verticalLabelPosition=bottom;verticalAlign=top;html=1;outlineConnect=0;fillColor=#F0FDF4;strokeColor=#15803D;strokeWidth=2;fontStyle=1;fontSize=12;" vertex="1" parent="1">
          <mxGeometry x="970" y="280" width="60" height="110" as="geometry" />
        </mxCell>

        <!-- Associations -->
        <mxCell id="assoc1" style="endArrow=none;html=1;rounded=0;strokeWidth=2;strokeColor=#1E3A8A;" edge="1" parent="1" source="act_shopper" target="uc1">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        
        <mxCell id="assoc2" style="endArrow=none;html=1;rounded=0;strokeWidth=2;strokeColor=#1E3A8A;" edge="1" parent="1" source="act_shopper" target="uc2">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        
        <mxCell id="assoc3" style="endArrow=none;html=1;rounded=0;strokeWidth=2;strokeColor=#1E3A8A;" edge="1" parent="1" source="act_shopper" target="uc3">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        
        <mxCell id="assoc4" style="endArrow=none;html=1;rounded=0;strokeWidth=2;strokeColor=#92400E;" edge="1" parent="1" source="act_artisan" target="uc6">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        
        <mxCell id="assoc5" style="endArrow=none;html=1;rounded=0;strokeWidth=2;strokeColor=#92400E;" edge="1" parent="1" source="act_artisan" target="uc7">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        
        <mxCell id="assoc6" style="endArrow=none;html=1;rounded=0;strokeWidth=2;strokeColor=#15803D;" edge="1" parent="1" source="act_gateway" target="uc4">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        
      </root>
    </mxGraphModel>
  </diagram>
</mxfile>"""
    drawio_diag_path = os.path.join(DIAGRAM_DIR, "artisan_marketplace_usecase.drawio")
    with open(drawio_diag_path, "w", encoding="utf-8") as f:
        f.write(drawio_xml)
    shutil.copyfile(drawio_diag_path, os.path.join(LAB1_DIR, "UseCase_Diagram.drawio"))
    print(f"  Generated Draw.io XML: {os.path.basename(drawio_diag_path)}")

# ==============================================================================
# 3. GENERATE USE-CASE FLOW SPECIFICATION (PDF, DOCX)
# ==============================================================================
def build_use_case_flow():
    print("\n--- Building Use-Case Flow Deliverables ---")

    # 3A. HTML -> PDF (Exactly 1 Page)
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>Use-Case Flow Specification - UC-03 Place Order</title>
<style>
  @page {{
    size: A4 portrait;
    margin: 7mm 11mm 7mm 11mm;
  }}
  * {{
    box-sizing: border-box;
    margin: 0;
    padding: 0;
  }}
  body {{
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif;
    color: #1f2937;
    background: #ffffff;
    line-height: 1.28;
    font-size: 8pt;
  }}
  .header-banner {{
    border: 1.5px solid #1e3a8a;
    border-radius: 5px;
    padding: 6px 12px;
    margin-bottom: 6px;
    background: linear-gradient(135deg, #f0f4f8 0%, #ffffff 100%);
    display: flex;
    justify-content: space-between;
    align-items: center;
  }}
  .header-banner h1 {{
    font-size: 11.5pt;
    color: #1e3a8a;
    font-weight: 700;
    letter-spacing: -0.2px;
  }}
  .header-banner h2 {{
    font-size: 8.5pt;
    color: #4b5563;
    font-weight: 600;
  }}
  .meta-grid {{
    width: 100%;
    border-collapse: collapse;
    margin-bottom: 6px;
  }}
  .meta-grid td {{
    padding: 3.5px 7px;
    border: 1px solid #cbd5e1;
    font-size: 7.6pt;
    vertical-align: middle;
  }}
  .meta-label {{
    background-color: #f1f5f9;
    font-weight: 700;
    color: #1e3a8a;
    width: 17%;
  }}
  .meta-value {{
    color: #0f172a;
    width: 33%;
  }}
  .section-title {{
    font-size: 8.2pt;
    font-weight: 700;
    color: #1e3a8a;
    text-transform: uppercase;
    letter-spacing: 0.4px;
    border-bottom: 1.5px solid #1e3a8a;
    padding-bottom: 2px;
    margin-top: 5px;
    margin-bottom: 4px;
  }}
  .conditions-box {{
    display: flex;
    gap: 8px;
    margin-bottom: 5px;
  }}
  .condition-col {{
    flex: 1;
    background: #f8fafc;
    border: 1px solid #e2e8f0;
    border-radius: 4px;
    padding: 5px 8px;
  }}
  .condition-col h3 {{
    font-size: 7.6pt;
    font-weight: 700;
    color: #1e3a8a;
    margin-bottom: 3px;
    text-transform: uppercase;
  }}
  ul {{
    padding-left: 14px;
  }}
  li {{
    margin-bottom: 2px;
    font-size: 7.4pt;
    color: #334155;
  }}
  .step-list {{
    margin-bottom: 5px;
  }}
  .step-row {{
    display: flex;
    padding: 2.2px 0;
    font-size: 7.5pt;
    border-bottom: 1px dotted #e2e8f0;
  }}
  .step-num {{
    font-weight: 700;
    color: #1e3a8a;
    width: 22px;
    flex-shrink: 0;
  }}
  .step-desc {{
    color: #1e293b;
  }}
  .highlight {{
    font-weight: 600;
    color: #1e3a8a;
  }}
  .alt-flow-box {{
    background: #fefce8;
    border: 1px solid #fef08a;
    border-left: 3.5px solid #ca8a04;
    border-radius: 4px;
    padding: 5px 8px;
    margin-bottom: 4px;
  }}
  .alt-flow-box.alt2 {{
    background: #eff6ff;
    border: 1px solid #bfdbfe;
    border-left: 3.5px solid #2563eb;
  }}
  .alt-title {{
    font-size: 7.7pt;
    font-weight: 700;
    color: #854d0e;
    margin-bottom: 2px;
  }}
  .alt-title.alt2 {{
    color: #1e40af;
  }}
  .alt-steps {{
    font-size: 7.3pt;
    color: #334155;
    line-height: 1.25;
  }}
  .alt-steps div {{
    margin-bottom: 1.5px;
  }}
  .footer {{
    border-top: 1px solid #cbd5e1;
    padding-top: 3px;
    margin-top: 4px;
    display: flex;
    justify-content: space-between;
    font-size: 6.8pt;
    color: #64748b;
  }}
</style>
</head>
<body>

<div class="header-banner">
  <div>
    <h1>PES UNIVERSITY &bull; DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING</h1>
    <h2>Software Engineering Lab (UE24CS252) &bull; Use-Case Flow Specification</h2>
  </div>
  <div style="text-align: right; font-size: 7.8pt;">
    <div><strong>Problem Statement #{PROBLEM_NUM}:</strong> {PROBLEM_TITLE}</div>
    <div><strong>Student:</strong> {STUDENT_NAME} ({STUDENT_SRN})</div>
  </div>
</div>

<table class="meta-grid">
  <tr>
    <td class="meta-label">Use Case ID &amp; Name</td>
    <td class="meta-value"><strong>UC-03: Place Order</strong></td>
    <td class="meta-label">System Scope</td>
    <td class="meta-value">Multi-Vendor Artisan Marketplace</td>
  </tr>
  <tr>
    <td class="meta-label">Primary Actor</td>
    <td class="meta-value"><strong>Shopper</strong></td>
    <td class="meta-label">Secondary Actors</td>
    <td class="meta-value"><strong>Payment Gateway</strong>, Artisan Vendor</td>
  </tr>
  <tr>
    <td class="meta-label">Trigger</td>
    <td class="meta-value">Shopper clicks "Proceed to Checkout" from cart</td>
    <td class="meta-label">Included / Extended UCs</td>
    <td class="meta-value">&lt;&lt;include&gt;&gt; UC-04 (Split Pay), &lt;&lt;extend&gt;&gt; UC-05 (Discount)</td>
  </tr>
</table>

<div class="conditions-box">
  <div class="condition-col">
    <h3>Preconditions</h3>
    <ul>
      <li>Shopper is authenticated with valid shipping address saved.</li>
      <li>Shopping cart contains items from &ge; 1 active artisan vendor storefronts.</li>
      <li>Item inventory quantities are verified available in catalog.</li>
      <li>Vendors have active, verified payout disbursement accounts.</li>
    </ul>
  </div>
  <div class="condition-col">
    <h3>Postconditions</h3>
    <ul>
      <li>Master order created and decomposed into vendor sub-orders.</li>
      <li>Consolidated cart payment authorized by Payment Gateway.</li>
      <li><strong>5% platform fee deducted; net payouts split to vendor escrows (FR-001).</strong></li>
      <li>Product inventory counts decremented; confirmation emails sent.</li>
    </ul>
  </div>
</div>

<div class="section-title">Main Success Scenario</div>
<div class="step-list">
  <div class="step-row">
    <div class="step-num">1.</div>
    <div class="step-desc">Shopper navigates to the checkout view from the multi-vendor shopping cart.</div>
  </div>
  <div class="step-row">
    <div class="step-num">2.</div>
    <div class="step-desc">System retrieves cart items across all artisan vendors, locks inventory quantities, and displays an itemized order summary.</div>
  </div>
  <div class="step-row">
    <div class="step-num">3.</div>
    <div class="step-desc">Shopper selects or confirms the delivery shipping address and recipient contact details.</div>
  </div>
  <div class="step-row">
    <div class="step-num">4.</div>
    <div class="step-desc">System computes shipping costs per artisan vendor, applicable taxes, and presents the grand total payable amount.</div>
  </div>
  <div class="step-row">
    <div class="step-num">5.</div>
    <div class="step-desc">Shopper selects payment method (Credit/Debit Card, UPI, Net Banking) and enters payment credentials.</div>
  </div>
  <div class="step-row">
    <div class="step-num">6.</div>
    <div class="step-desc">System invokes <span class="highlight">&lt;&lt;include&gt;&gt; UC-04: Process Split Payment</span>, routing encrypted transaction payload (TLS 1.3) to the <span class="highlight">Payment Gateway</span>.</div>
  </div>
  <div class="step-row">
    <div class="step-num">7.</div>
    <div class="step-desc">Payment Gateway authorizes the full consolidated cart amount and returns a success authorization token.</div>
  </div>
  <div class="step-row">
    <div class="step-num">8.</div>
    <div class="step-desc">System automatically <span class="highlight">deducts the 5% platform commission</span> and allocates remaining item earnings to respective artisan vendor payout accounts (adhering to <span class="highlight">FR-001</span>).</div>
  </div>
  <div class="step-row">
    <div class="step-num">9.</div>
    <div class="step-desc">System decrements catalog inventory quantities, decomposes master order into isolated vendor sub-orders, and schedules disbursement settlements.</div>
  </div>
  <div class="step-row">
    <div class="step-num">10.</div>
    <div class="step-desc">System displays an order confirmation screen with order ID and transmits confirmation receipts to the shopper and alert notifications to respective artisan vendors. Use case ends successfully.</div>
  </div>
</div>

<div class="section-title">Alternate Flows</div>
<div class="alt-flow-box">
  <div class="alt-title">6a. Alternate Flow: Payment Authorization Declined (at Step 6 / 7)</div>
  <div class="alt-steps">
    <div><strong>6a1.</strong> Payment Gateway returns an authorization failure code (e.g., insufficient funds, bank timeout, expired card).</div>
    <div><strong>6a2.</strong> System logs the transaction attempt, flags order as 'Payment Pending', and presents a clear error message to the shopper without releasing cart items.</div>
    <div><strong>6a3.</strong> System prompts the shopper to select an alternate payment method or re-enter valid credentials.</div>
    <div><strong>6a4.</strong> If shopper successfully resubmits payment within 3 attempts, flow resumes at <strong>Step 7</strong>.</div>
    <div><strong>6a5.</strong> If all payment attempts fail or shopper abandons session, locked inventory is released back to artisan catalogs after 15 minutes and the checkout process terminates.</div>
  </div>
</div>

<div class="alt-flow-box alt2">
  <div class="alt-title alt2">4a. Alternate / Extension Flow: Apply Promotional Discount Code (&lt;&lt;extend&gt;&gt; UC-05, at Step 4)</div>
  <div class="alt-steps">
    <div><strong>4a1.</strong> Shopper enters a promotional coupon code in the checkout discount field and clicks 'Apply'.</div>
    <div><strong>4a2.</strong> System validates coupon eligibility against vendor items and cart criteria. If valid, system recalculates the discounted total and updates vendor payout bases; flow resumes at <strong>Step 5</strong>.</div>
    <div><strong>4a3.</strong> If coupon is invalid or expired, system displays an inline warning and retains the original cart total without interrupting checkout.</div>
  </div>
</div>

<div class="footer">
  <span>Software Engineering Lab (UE24CS252) &bull; PES University Dept. of CSE</span>
  <span>{STUDENT_SRN} &bull; {STUDENT_NAME}</span>
  <span>Problem Statement #{PROBLEM_NUM}: {PROBLEM_TITLE} &bull; Page 1 of 1</span>
</div>

</body>
</html>
"""
    html_path = os.path.join(DOCS_DIR, "03_UseCase_Flow_Specification.html")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html)

    pdf_docs_path = os.path.join(DOCS_DIR, "03_UseCase_Flow_Specification.pdf")
    pdf_root_path = os.path.join(LAB1_DIR, "Use_Case_Flow_Specification.pdf")
    print_html_to_pdf(html_path, pdf_docs_path, landscape=False)
    shutil.copyfile(pdf_docs_path, pdf_root_path)
    shutil.copyfile(pdf_docs_path, os.path.join(LAB1_DIR, "UseCase_Flow_Specification.pdf"))

    # 3B. DOCX (Word Document)
    docx_path = os.path.join(DOCS_DIR, "03_UseCase_Flow_Specification.docx")
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.orientation = docx.enum.section.WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

    title_p = doc.add_paragraph()
    r1 = title_p.add_run("PES UNIVERSITY • DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING\n")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    r2 = title_p.add_run(f"Use-Case Flow Specification • UC-03: Place Order\nProblem Statement #{PROBLEM_NUM}: {PROBLEM_TITLE} | Student: {STUDENT_NAME} ({STUDENT_SRN})\n")
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    meta_tbl = doc.add_table(rows=3, cols=4)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        [("Use Case ID & Name", True), ("UC-03: Place Order", False), ("Scope", True), ("Multi-Vendor Artisan Marketplace", False)],
        [("Primary Actor", True), ("Shopper", False), ("Secondary Actors", True), ("Payment Gateway, Artisan Vendor", False)],
        [("Trigger", True), ("Clicks 'Proceed to Checkout'", False), ("Stereotypes", True), ("«include» UC-04, «extend» UC-05", False)]
    ]
    for r_idx, row in enumerate(meta_data):
        for c_idx, (text, is_hdr) in enumerate(row):
            cell = meta_tbl.cell(r_idx, c_idx)
            cell.text = text
            if is_hdr:
                shd = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shd)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)
                    r.font.name = "Segoe UI"
                    if is_hdr:
                        r.bold = True
                        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    doc.add_paragraph()

    def add_sec(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        return p

    add_sec("1. Preconditions & Postconditions")
    pre_p = doc.add_paragraph()
    pre_p.paragraph_format.space_after = Pt(2)
    pre_p.add_run("Preconditions:\n").bold = True
    pre_p.add_run("• Shopper is authenticated with valid delivery address saved.\n• Cart contains items from >= 1 active artisan storefronts.\n• Real-time stock counts verified available in catalog.\n• Vendors have active, verified payout disbursement accounts.")
    
    post_p = doc.add_paragraph()
    post_p.paragraph_format.space_after = Pt(2)
    post_p.add_run("Postconditions:\n").bold = True
    post_p.add_run("• Master order created and decomposed into vendor sub-orders.\n• Consolidated payment authorized by Payment Gateway.\n• 5% platform fee deducted; net payouts split to vendor escrows (FR-001).\n• Product stock decremented; confirmation receipts transmitted.")

    add_sec("2. Main Success Scenario")
    steps = [
        "Shopper navigates to the checkout view from the multi-vendor shopping cart.",
        "System retrieves cart items across all artisan storefronts, locks inventory, and presents order summary.",
        "Shopper selects or confirms the delivery shipping address and contact details.",
        "System computes shipping costs per artisan vendor, applicable taxes, and presents total payable amount.",
        "Shopper selects payment method (Credit/Debit Card, UPI, Net Banking) and enters payment credentials.",
        "System invokes «include» UC-04: Process Split Payment, routing encrypted payload (TLS 1.3) to Payment Gateway.",
        "Payment Gateway authorizes the full consolidated cart amount and returns authorization token.",
        "System automatically deducts the 5% platform commission and allocates net payouts to respective vendor accounts (FR-001).",
        "System decrements catalog inventory quantities, decomposes master order into isolated vendor sub-orders, and schedules disbursement settlements.",
        "System displays order confirmation screen with order ID and dispatches confirmation receipts. Use case ends successfully."
    ]
    for idx, s in enumerate(steps, 1):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(1)
        sp.paragraph_format.space_after = Pt(1)
        sp.add_run(f"{idx}. ").bold = True
        sp.add_run(s)

    add_sec("3. Alternate & Extension Flows")
    alt1 = doc.add_paragraph()
    alt1.paragraph_format.space_after = Pt(2)
    alt1.add_run("6a. Alternate Flow: Payment Authorization Declined (at Step 6/7)\n").bold = True
    alt1.add_run("• 6a1. Payment Gateway returns an authorization failure code (insufficient funds, timeout, expired card).\n• 6a2. System logs attempt, flags order as 'Payment Pending', and presents error message without releasing cart items.\n• 6a3. System prompts shopper to select an alternate payment method (up to 3 attempts).\n• 6a4. If successful, flow resumes at Step 7; if abandoned, locked inventory is released back after 15 minutes and checkout terminates.")

    alt2 = doc.add_paragraph()
    alt2.paragraph_format.space_after = Pt(2)
    alt2.add_run("4a. Extension Flow: Apply Promotional Discount Code («extend» UC-05, at Step 4)\n").bold = True
    alt2.add_run("• 4a1. Shopper enters a promotional coupon code in checkout discount field and clicks 'Apply'.\n• 4a2. System validates coupon eligibility. If valid, system recalculates discounted total and updates vendor payout bases; flow resumes at Step 5.\n• 4a3. If coupon is invalid or expired, system displays warning and retains original cart total.")

    doc.save(docx_path)
    shutil.copyfile(docx_path, os.path.join(LAB1_DIR, "UseCase_Flow_PlaceOrder.docx"))
    shutil.copyfile(docx_path, os.path.join(LAB1_DIR, "UseCase_Flow_Specification.docx"))
    print(f"  Generated DOCX: {os.path.basename(docx_path)}")

# ==============================================================================
# 4. GENERATE LAB1/README.MD AND ROOT README.MD
# ==============================================================================
def build_readmes():
    print("\n--- Building Comprehensive README Documentation ---")

    # LAB1/README.md
    lab1_readme = f"""# Lab 1: Requirements Engineering & UML Use-Case Modelling

**Institution:** PES University • Department of Computer Science & Engineering  
**Course:** Software Engineering (UE24CS252)  
**Problem Statement #{PROBLEM_NUM}:** {PROBLEM_TITLE}  
**Domain:** {DOMAIN}  
**Student:** {STUDENT_NAME} ({STUDENT_SRN})  

---

## 1. Problem Context & Overview

The **Multi-Vendor Artisan E-Commerce Marketplace** is an online platform that empowers independent craftspeople and artisans to establish digital storefronts, showcase handcrafted product catalogs with rich media, process customer orders, and receive automated split payouts with platform commission deductions.

### Target Stakeholders & Actors
1. **Shopper (Primary Human Actor):** Discovers unique handcrafted items, aggregates goods from multiple distinct artisans in a single unified cart, applies coupons, and executes consolidated checkout orders.
2. **Artisan Vendor (Primary Human Actor):** Manages artisan storefront profile, lists products with craft attributes and stock quantities, monitors incoming sub-orders, and fulfills independent shipments with carrier tracking.
3. **Payment Gateway (Secondary External System Actor):** Securely authorizes multi-party credit card/UPI transactions, validates anti-fraud tokens, and routes settlement disbursement instructions.
4. **Platform Administrator (Secondary Actor):** Supervises platform-wide commission rules (5% deduction), monitors financial audit reconciliations, and handles vendor dispute arbitrations.

---

## 2. Requirements Specification Table

*Complete Requirements Table conforming to the Lab 1 Handout specifications: exactly 5 Functional Requirements (FR-001 to FR-005) and 2 Non-Functional Requirements (NFR-001 & NFR-002).*

| Req ID | Type | Description | Priority | Acceptance Criteria | Rationale | Comments |
| :--- | :--- | :--- | :---: | :--- | :--- | :--- |
| **FR-001** | Functional | The system shall split customer cart payments at checkout, allocating respective item earnings to multiple independent vendor accounts after deducting a 5% platform fee. | **High** | **Pass:** Split payout calculations balance to total cart value.<br>**Fail:** Payout calculation discrepancies. | Ensures accurate, automated financial distribution to independent artisans while securing platform transaction revenue. *(Given in PS #31)* | Core monetization logic; math must strictly balance to 100% of payable order total. |
| **FR-002** | Functional | The system shall allow artisan vendors to set up and manage an independent storefront profile, including artisan bio, shop policies, and linked payout disbursement account details. | **High** | **Pass:** Storefront profile updates and verified payout details are successfully saved and rendered publicly on the artisan's storefront page.<br>**Fail:** Missing mandatory profile fields or unvalidated payout details prevent storefront publishing. | Enables independent craftspeople to establish their distinct brand identity and receive automated financial payouts. | Peer critique: Payout bank credentials must undergo automated format validation prior to store activation. |
| **FR-003** | Functional | The system shall allow artisan vendors to create, update, and manage handcrafted product listings with titles, descriptions, pricing, inventory quantities, and high-resolution product media. | **High** | **Pass:** New or updated product listings appear in the artisan's catalog and marketplace search within 5 seconds with accurate stock counts.<br>**Fail:** Listings with invalid pricing (≤ 0) or missing mandatory attributes fail validation and are rejected. | Empowers artisans to maintain an up-to-date catalog of handcrafted items and prevent overselling through inventory synchronization. | Real-time stock reservation prevents race conditions during concurrent shopper checkouts. |
| **FR-004** | Functional | The system shall allow shoppers to add handcrafted items from multiple independent artisan vendors into a unified cart and execute a single consolidated checkout order. | **High** | **Pass:** Cart aggregates items across distinct vendors, calculates itemized totals, applies applicable coupon discounts, and initiates payment authorization.<br>**Fail:** Cart fails to itemize multi-vendor items or calculation mismatch occurs between item sums and total payable amount. | Provides a seamless purchasing experience for shoppers buying from multiple artisans simultaneously without requiring separate checkout transactions. | Master order automatically decomposes into vendor sub-orders upon successful authorization. |
| **FR-005** | Functional | The system shall notify artisan vendors of confirmed customer orders containing their respective products and allow vendors to update fulfillment status (Processing, Dispatched, Delivered) with carrier tracking details. | **Medium** | **Pass:** Vendor dashboard immediately displays incoming sub-orders upon successful checkout and status transitions trigger automated notifications to the shopper.<br>**Fail:** Order details fail to isolate vendor-specific items or status updates fail to persist. | Enables independent artisans to fulfill customer orders independently while keeping shoppers informed of dispatch status. | Status transitions dispatch transactional email/push notifications containing carrier tracking links. |
| **NFR-001** | Performance & Security | The product catalog must support high-resolution image rendering with CDN caching delivering load times < 500 ms. | **High** | **Pass:** Benchmarking tests confirm target latency (< 500 ms) and security standards under simulated peak load.<br>**Fail:** Catalog page load latency ≥ 500 ms or CDN caching failure under peak load. | Ensures fast page loading for image-heavy handcrafted goods, preserving user engagement and minimizing bounce rates. *(Given in PS #31)* | Optimized WebP/AVIF imagery served via multi-region edge CDN caches; validated via Lighthouse. |
| **NFR-002** | Security & Compliance | The system shall encrypt all sensitive financial transactions and payout account data in transit using TLS 1.3 and at rest using AES-256 encryption, adhering to PCI-DSS Level 1 compliance standards. | **High** | **Pass:** Automated vulnerability scans and penetration audits confirm 100% encryption coverage for payment payloads with zero plain-text storage of payment credentials.<br>**Fail:** Any unencrypted transmission or unmasked storage of sensitive banking/card data detected. | Protects shoppers' financial credentials and vendors' banking payout information against unauthorized access, data breaches, and regulatory non-compliance. | Double-entry append-only transaction ledger ensures 99.999% audit reconciliation accuracy. |

---

## 3. UML Use-Case Diagram

![UML Use-Case Diagram](docs/02_UseCase_Diagram.png)

*The vector diagram is submitted as [`docs/02_UseCase_Diagram.pdf`](docs/02_UseCase_Diagram.pdf). The diagram can be viewed and edited in [draw.io / diagrams.net](https://app.diagrams.net) via [`diagram/artisan_marketplace_usecase.drawio`](diagram/artisan_marketplace_usecase.drawio).*

### Actor & Use Case Traceability

| Use Case ID | Use Case Title | Primary Actor / Relationship | Traces to Requirement |
| :--- | :--- | :--- | :--- |
| **UC-01** | Browse & Search Catalog | Shopper | FR-003, NFR-001 |
| **UC-02** | Manage Shopping Cart | Shopper | FR-004 |
| **UC-03** | Place Order [Multi-Vendor Checkout] | Shopper | FR-001, FR-004, NFR-002 |
| **UC-04** | Process Split Payment | Payment Gateway; *`«include»` by UC-03* | FR-001, NFR-002 |
| **UC-05** | Apply Discount Code | Shopper; *`«extend»` to UC-03* | FR-004 |
| **UC-06** | Manage Storefront & Catalog | Artisan Vendor | FR-002, FR-003, NFR-001 |
| **UC-07** | Fulfill Order | Artisan Vendor | FR-005 |

### UML Stereotype Justifications
1. **`«include»` Relationship (UC-03 ➔ UC-04):**  
   * **Source:** `UC-03: Place Order` | **Target:** `UC-04: Process Split Payment`  
   * **Semantics:** Placing an order in a multi-vendor marketplace *unconditionally requires* the split payment mechanism to execute. The system cannot complete checkout without calculating and deducting the 5% platform commission and allocating item revenues to respective vendors. Therefore, `UC-04` is mandatory and included in `UC-03`.
2. **`«extend»` Relationship (UC-05 ➔ UC-03):**  
   * **Source:** `UC-05: Apply Discount Code` | **Target:** `UC-03: Place Order`  
   * **Semantics:** Applying a coupon code is *optional, conditional behavior* that occurs only if the shopper possesses and enters an eligible promotional code. The base checkout workflow operates successfully without a discount; hence, `UC-05` extends `UC-03` at the extension point **"Order Total Calculation"**.

---

## 4. Use-Case Flow Specification

### UC-03: Place Order (Multi-Vendor Checkout with Split Payment)

* **Scope:** Multi-Vendor Artisan E-Commerce Marketplace
* **Primary Actor:** Shopper
* **Secondary Actors:** Payment Gateway (external), Artisan Vendor
* **Trigger:** Shopper selects "Proceed to Checkout" from the active shopping cart view.
* **Traced Requirements:** FR-001, FR-004, NFR-002
* **Relationships:** `«include»` UC-04 Process Split Payment; extended by `«extend»` UC-05 Apply Discount Code

#### Preconditions
1. Shopper is authenticated and has a valid shipping delivery address on file.
2. Shopping cart contains at least one handcrafted item from one or more active artisan storefronts.
3. Item inventory quantities are verified available in the real-time catalog.
4. Each artisan vendor represented in the cart has a verified, active payout account on record.

#### Postconditions
1. A master order record is created and persisted with status `Confirmed`.
2. The master order is decomposed into isolated sub-orders corresponding to each distinct artisan vendor.
3. Full order payment is authorized and captured via the Payment Gateway.
4. The 5% platform fee is deducted and transferred to the marketplace revenue account; net payouts are allocated to vendor escrow ledgers (**FR-001**).
5. Catalog inventory quantities for purchased items are decremented.
6. Order confirmation receipts are dispatched to the shopper and fulfillment alerts sent to artisan vendors.
7. *Failure Guarantee:* In the event of payment or inventory failure, no charges are settled, inventory holds are released, and the cart remains intact.

#### Main Success Scenario
1. Shopper navigates to the checkout view from the multi-vendor shopping cart.
2. System retrieves cart items across all artisan vendors, places a temporary 15-minute hold on inventory counts, and displays an itemized order summary.
3. Shopper selects or confirms the delivery shipping address and contact phone number.
4. System computes shipping costs per artisan vendor, applicable taxes, and presents the grand total payable amount.
5. Shopper selects payment method (Credit/Debit Card, UPI, Net Banking) and enters payment credentials.
6. System invokes **`«include»` UC-04: Process Split Payment**, encrypting transaction payload via TLS 1.3 and dispatching it to the **Payment Gateway**.
7. Payment Gateway authorizes the full consolidated cart amount and returns a success authorization token.
8. System automatically **deducts the 5% platform commission** and allocates remaining item earnings to respective artisan vendor payout accounts (adhering to **FR-001**).
9. System decrements catalog inventory quantities, decomposes master order into isolated vendor sub-orders, and schedules disbursement settlements.
10. System displays an order confirmation screen with unique order IDs and transmits confirmation receipts to the shopper and alert notifications to respective artisan vendors. Use case ends successfully.

#### Alternate Flows
* **6a. Alternate Flow: Payment Authorization Declined (at Step 6 / 7)**
  * **6a1.** Payment Gateway returns an authorization failure code (e.g., insufficient funds, bank timeout, expired card).
  * **6a2.** System logs the transaction attempt, flags order as 'Payment Pending', and presents a clear error message to the shopper without releasing cart items.
  * **6a3.** System prompts the shopper to select an alternate payment method or re-enter valid credentials.
  * **6a4.** If shopper successfully resubmits payment within 3 attempts, flow resumes at **Step 7**.
  * **6a5.** If all payment attempts fail or shopper abandons session, locked inventory is released back to artisan catalogs after 15 minutes and the checkout process terminates.
* **4a. Alternate / Extension Flow: Apply Promotional Discount Code (`«extend»` UC-05, at Step 4)**
  * **4a1.** Shopper enters a promotional coupon code in the checkout discount field and clicks 'Apply'.
  * **4a2.** System validates coupon eligibility against vendor items and cart criteria. If valid, system recalculates the discounted total and updates vendor payout bases; flow resumes at **Step 5**.
  * **4a3.** If coupon is invalid or expired, system displays an inline warning and retains the original cart total without interrupting checkout.

---

## 5. Lab Deliverables Directory

All files have been systematically organized into formal directories:

```
LAB1/
├── README.md                                   # Comprehensive Lab 1 documentation (this document)
├── Requirements_Table.pdf                      # Requirements Table (A4 Landscape PDF)
├── Use_Case_Diagram.pdf                        # UML Use-Case Diagram (A4 Landscape PDF)
├── Use_Case_Flow_Specification.pdf             # Use-Case Flow Specification (Exactly 1-page A4 Portrait PDF)
├── Use_Case_Diagram.svg                        # Standalone vector SVG diagram
├── Use_Case_Diagram.drawio                     # Editable Draw.io diagram source
├── use_case_diagram.png                        # High-resolution PNG preview
│
├── docs/                                       # Formatted submission documents
│   ├── 01_Requirements_Table.pdf              # Publication-grade styled PDF
│   ├── 01_Requirements_Table.docx             # Editable Microsoft Word document
│   ├── 01_Requirements_Table.xlsx             # Formatted Microsoft Excel spreadsheet
│   ├── 02_UseCase_Diagram.pdf                 # Vector PDF diagram
│   ├── 02_UseCase_Diagram.png                 # High-resolution 1080x700 PNG image
│   ├── 03_UseCase_Flow_Specification.pdf      # Exactly 1-page formal flow PDF
│   └── 03_UseCase_Flow_Specification.docx     # Editable Microsoft Word document
│
└── diagram/                                   # Diagram design sources
    ├── artisan_marketplace_usecase.drawio     # Editable XML source for diagrams.net / draw.io
    └── artisan_marketplace_usecase.svg        # Scalable Vector Graphics source
```
"""
    lab1_readme_path = os.path.join(LAB1_DIR, "README.md")
    with open(lab1_readme_path, "w", encoding="utf-8") as f:
        f.write(lab1_readme)
    print(f"  Generated: LAB1/README.md")

    # Root README.md
    root_readme = f"""# Software Engineering Lab Submissions (UE24CS252)

**Student Name:** {STUDENT_NAME}  
**SRN:** {STUDENT_SRN}  
**Department:** Computer Science & Engineering  
**Institution:** PES University  

---

## Lab Index

| Lab | Topic | Problem Statement | Domain | Status | Deliverables |
| :---: | :--- | :--- | :--- | :---: | :--- |
| **Lab 1** | Requirements Engineering & UML Use-Case Modelling | **#{PROBLEM_NUM}: {PROBLEM_TITLE}** | {DOMAIN} | **Completed** | [Lab 1 Folder](LAB1/) &bull; [Docs](LAB1/docs/) |

---

## Lab 1: Multi-Vendor Artisan E-Commerce Marketplace

An online marketplace enabling independent craftspeople to set up storefronts, manage product catalogs, receive orders, and receive automated split payouts with platform commission deductions.

### Quick Deliverable Links
- 📋 **Requirements Table:** [`01_Requirements_Table.pdf`](LAB1/docs/01_Requirements_Table.pdf) · [`01_Requirements_Table.docx`](LAB1/docs/01_Requirements_Table.docx) · [`01_Requirements_Table.xlsx`](LAB1/docs/01_Requirements_Table.xlsx)
- 📊 **UML Use-Case Diagram:** [`02_UseCase_Diagram.pdf`](LAB1/docs/02_UseCase_Diagram.pdf) · [`02_UseCase_Diagram.png`](LAB1/docs/02_UseCase_Diagram.png) · [Editable `.drawio`](LAB1/diagram/artisan_marketplace_usecase.drawio)
- 📄 **Use-Case Flow Spec:** [`03_UseCase_Flow_Specification.pdf`](LAB1/docs/03_UseCase_Flow_Specification.pdf) · [`03_UseCase_Flow_Specification.docx`](LAB1/docs/03_UseCase_Flow_Specification.docx)

---

### UML Use-Case Diagram

![UML Use-Case Diagram](LAB1/docs/02_UseCase_Diagram.png)

*To edit the diagram, open [`LAB1/diagram/artisan_marketplace_usecase.drawio`](LAB1/diagram/artisan_marketplace_usecase.drawio) in [app.diagrams.net (Draw.io)](https://app.diagrams.net).*

---

### Requirements Summary Table

| Req ID | Type | Description | Priority | Acceptance Criteria | Rationale |
| :--- | :--- | :--- | :---: | :--- | :--- |
| **FR-001** | Functional | The system shall split customer cart payments at checkout, allocating respective item earnings to multiple independent vendor accounts after deducting a 5% platform fee. | **High** | **Pass:** Split payout calculations balance to total cart value.<br>**Fail:** Payout calculation discrepancies. | Ensures accurate, automated financial distribution to independent artisans while securing platform transaction revenue. *(Given in PS #31)* |
| **FR-002** | Functional | The system shall allow artisan vendors to set up and manage an independent storefront profile, including artisan bio, shop policies, and linked payout disbursement account details. | **High** | **Pass:** Storefront profile updates and verified payout details are successfully saved and rendered publicly.<br>**Fail:** Missing mandatory fields or unvalidated payout details prevent publishing. | Enables independent craftspeople to establish their distinct brand identity and receive automated financial payouts. |
| **FR-003** | Functional | The system shall allow artisan vendors to create, update, and manage handcrafted product listings with titles, descriptions, pricing, inventory quantities, and high-resolution product media. | **High** | **Pass:** New/updated product listings appear in search within 5 seconds with accurate stock counts.<br>**Fail:** Listings with invalid pricing (≤ 0) or missing mandatory attributes fail validation. | Empowers artisans to maintain an up-to-date catalog of handcrafted items and prevent overselling. |
| **FR-004** | Functional | The system shall allow shoppers to add handcrafted items from multiple independent artisan vendors into a unified cart and execute a single consolidated checkout order. | **High** | **Pass:** Cart aggregates items across distinct vendors, calculates itemized totals, applies coupons, and initiates payment.<br>**Fail:** Cart fails to itemize multi-vendor items or calculation mismatch occurs. | Provides a seamless purchasing experience for shoppers buying from multiple artisans simultaneously. |
| **FR-005** | Functional | The system shall notify artisan vendors of confirmed customer orders containing their respective products and allow vendors to update fulfillment status (Processing, Dispatched, Delivered) with carrier tracking details. | **Medium** | **Pass:** Vendor dashboard immediately displays incoming sub-orders and status transitions trigger automated notifications.<br>**Fail:** Order details fail to isolate vendor-specific items or status updates fail to persist. | Enables independent artisans to fulfill customer orders independently while keeping shoppers informed. |
| **NFR-001** | Performance & Security | The product catalog must support high-resolution image rendering with CDN caching delivering load times < 500 ms. | **High** | **Pass:** Benchmarking tests confirm target latency (< 500 ms) and security standards under simulated peak load.<br>**Fail:** Catalog page load latency ≥ 500 ms or CDN caching failure under peak load. | Ensures fast page loading for image-heavy handcrafted goods, preserving user engagement and minimizing bounce rates. *(Given in PS #31)* |
| **NFR-002** | Security & Compliance | The system shall encrypt all sensitive financial transactions and payout account data in transit using TLS 1.3 and at rest using AES-256 encryption, adhering to PCI-DSS Level 1 compliance standards. | **High** | **Pass:** Automated audits confirm 100% encryption coverage for payment payloads with zero plain-text storage of credentials.<br>**Fail:** Any unencrypted transmission or unmasked storage of sensitive banking/card data detected. | Protects shoppers' financial credentials and vendors' banking payout information against unauthorized access and data breaches. |

---

### Core Use Case Flow: UC-03 Place Order (Multi-Vendor Checkout)

* **Primary Actor:** Shopper &bull; **Secondary Actors:** Payment Gateway, Artisan Vendor
* **Traces to:** FR-001, FR-004, NFR-002 &bull; **Stereotypes:** `«include»` UC-04 Process Split Payment, `«extend»` UC-05 Apply Discount Code

#### Main Success Scenario
1. Shopper navigates to the checkout view from the multi-vendor shopping cart.
2. System retrieves cart items across all artisan vendors, locks inventory quantities, and displays an itemized order summary.
3. Shopper selects or confirms the delivery shipping address and recipient contact details.
4. System computes shipping costs per artisan vendor, applicable taxes, and presents the grand total payable amount.
5. Shopper selects payment method (Credit/Debit Card, UPI, Net Banking) and enters payment credentials.
6. System invokes **`«include»` UC-04: Process Split Payment**, routing encrypted transaction payload (TLS 1.3) to the **Payment Gateway**.
7. Payment Gateway authorizes the full consolidated cart amount and returns a success authorization token.
8. System automatically **deducts the 5% platform commission** and allocates remaining item earnings to respective artisan vendor payout accounts (adhering to **FR-001**).
9. System decrements catalog inventory quantities, decomposes master order into isolated vendor sub-orders, and schedules disbursement settlements.
10. System displays an order confirmation screen with order ID and transmits confirmation receipts to the shopper and alert notifications to respective artisan vendors. Use case ends successfully.

---

For complete detailed specifications, visit the [`LAB1/` directory](LAB1/).
"""
    root_readme_path = os.path.join(BASE_DIR, "README.md")
    with open(root_readme_path, "w", encoding="utf-8") as f:
        f.write(root_readme)
    print(f"  Generated: Root README.md")

# ==============================================================================
# MAIN EXECUTION
# ==============================================================================
if __name__ == "__main__":
    print("=================================================================")
    print(f" GENERATING SE LAB 1 DELIVERABLES FOR {STUDENT_NAME} ({STUDENT_SRN})")
    print(f" Problem Statement #{PROBLEM_NUM}: {PROBLEM_TITLE}")
    print("=================================================================")

    build_requirements_table()
    build_use_case_diagram()
    build_use_case_flow()
    build_readmes()

    print("\n=================================================================")
    print(" VERIFYING GENERATED PDF PAGE COUNTS")
    print("=================================================================")
    pdfs_to_check = [
        ("01_Requirements_Table.pdf", os.path.join(DOCS_DIR, "01_Requirements_Table.pdf")),
        ("02_UseCase_Diagram.pdf", os.path.join(DOCS_DIR, "02_UseCase_Diagram.pdf")),
        ("03_UseCase_Flow_Specification.pdf", os.path.join(DOCS_DIR, "03_UseCase_Flow_Specification.pdf")),
    ]
    all_passed = True
    for name, path in pdfs_to_check:
        pages = get_pdf_page_count(path)
        print(f"  {name}: {pages} page(s)")
        if pages != 1:
            print(f"  [WARNING] {name} is {pages} pages, expected 1 page!")
            all_passed = False

    if all_passed:
        print("\nSUCCESS: All PDF deliverables are exactly 1 page and formatted perfectly!")
    else:
        print("\nATTENTION: Some deliverables exceeded 1 page.")
    print("=================================================================")
